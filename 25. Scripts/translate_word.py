import os
from docx import Document
from deep_translator import GoogleTranslator
 
def traduzir_texto(texto, source='auto', target='pt'):
    """Traduz o texto mantendo o conteúdo se for vazio ou curto demais."""
    if not texto or texto.strip() == "" or len(texto.strip()) < 1:
        return texto if texto is not None else ""
    try:
        traduzido = GoogleTranslator(source=source, target=target).translate(texto)
        return traduzido if traduzido is not None else texto
    except Exception as e:
        print(f"Erro na tradução: {e}")
        return texto
 
def traduzir_word(arquivo_entrada, idioma_destino='pt'):
    doc = Document(arquivo_entrada)
    print(f"Iniciando tradução de: {arquivo_entrada}...")
 
    # Traduzindo Parágrafos
    for p in doc.paragraphs:
        for run in p.runs:
            if run.text.strip():
                run.text = traduzir_texto(run.text, target=idioma_destino)
 
    # Traduzindo Tabelas
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    for run in paragrafo.runs:
                        if run.text.strip():
                            run.text = traduzir_texto(run.text, target=idioma_destino)
 
    arquivo_saida = f"traduzido_{arquivo_entrada}"
    doc.save(arquivo_saida)
    print(f"Sucesso! Arquivo salvo como: {arquivo_saida}")
 
# --- USO ---
# Altere o nome do arquivo e o idioma de destino ('en', 'es', 'fr', etc.)
# Nota: O caminho deve ser relativo ao diretório de execução ou absoluto.
import os
arquivo = os.path.join("WORD", "BA 89503126_000100 _Innopal_EN_02.docx")
traduzir_word(arquivo, idioma_destino='pt')
